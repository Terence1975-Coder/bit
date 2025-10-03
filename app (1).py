import io
import os
import subprocess
import tempfile
from datetime import timedelta

import streamlit as st
from docx import Document

# Try portable ffmpeg via imageio-ffmpeg (no system PATH required)
FFMPEG_PATH = None
try:
    import imageio_ffmpeg
    FFMPEG_PATH = imageio_ffmpeg.get_ffmpeg_exe()
except Exception:
    FFMPEG_PATH = None  # We'll fall back to system "ffmpeg" if available

# Supported file types
SUPPORTED_TYPES = ["mp4", "mp3", "wav", "m4a", "mov", "mkv", "avi"]

# -------- Utils --------
def extract_audio_to_wav(in_bytes: bytes, out_path: str):
    """
    Convert any audio/video bytes to 16kHz mono WAV using ffmpeg.
    Uses portable imageio-ffmpeg binary if present; otherwise relies on system ffmpeg on PATH.
    """
    in_file = tempfile.NamedTemporaryFile(delete=False, suffix=".input")
    in_file.write(in_bytes)
    in_file.flush()
    in_file.close()

    ffmpeg_bin = FFMPEG_PATH or "ffmpeg"
    cmd = [ffmpeg_bin, "-y", "-i", in_file.name, "-ac", "1", "-ar", "16000", "-f", "wav", out_path]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except FileNotFoundError:
        raise RuntimeError(
            "FFmpeg binary not found. Install system FFmpeg (e.g. 'winget install Gyan.FFmpeg') "
            "or install Python package 'imageio-ffmpeg' so a portable binary is available."
        )
    finally:
        try: os.unlink(in_file.name)
        except OSError: pass

def format_timestamp(seconds: float) -> str:
    td = timedelta(seconds=float(seconds))
    hours, remainder = divmod(td.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    milliseconds = int(td.microseconds / 1000)
    hours += td.days * 24
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"

def segments_to_plaintext(segments):
    return "\n".join(seg.text.strip() for seg in segments)

def segments_to_srt(segments):
    lines = []
    for i, seg in enumerate(segments, start=1):
        start = format_timestamp(seg.start)
        end = format_timestamp(seg.end)
        text = seg.text.strip()
        lines.append(f"{i}\n{start} --> {end}\n{text}\n")
    return "\n".join(lines)

def write_docx(segments) -> bytes:
    doc = Document()
    doc.add_heading("Transcription", 0)
    for seg in segments:
        p = doc.add_paragraph()
        p.add_run(
            f"[{format_timestamp(seg.start).replace(',', '.')} - "
            f"{format_timestamp(seg.end).replace(',', '.')}] "
        ).bold = True
        p.add_run(seg.text.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# -------- Model loader with safe fallbacks --------
@st.cache_resource(show_spinner=False)
def load_model(model_size: str, compute_type: str):
    from faster_whisper import WhisperModel
    tried = set()
    # Try user choice first, then CPU-safe fallbacks
    candidates = [compute_type, "int8", "int8_float32", "float32"]
    last_err = None
    for ct in candidates:
        if ct in tried:
            continue
        try:
            return WhisperModel(model_size, compute_type=ct)
        except Exception as e:
            last_err = e
            tried.add(ct)
    raise last_err

# -------- Streamlit UI --------
st.set_page_config(page_title="MP4 → Word Transcriber", page_icon="📝", layout="centered")
st.title("📝 MP4 → Word Transcriber")

with st.sidebar:
    st.header("Settings")
    model_size = st.selectbox(
        "Model",
        ["small.en", "small", "medium.en", "medium", "large-v2"],
        index=0,
        help="*.en = English-only (faster). Non-.en = multilingual."
    )
    compute_type = st.selectbox(
        "Compute type",
        ["float32", "int8", "int8_float32", "float16"],
        index=0,
        help="On CPU use float32 / int8; on GPU use float16 if available."
    )
    vad_filter = st.toggle("VAD speech detection", value=True, help="Helps skip long silences/noise.")
    beam_size = st.slider("Beam size", 1, 10, 5)

uploaded = st.file_uploader("Upload audio/video file", type=SUPPORTED_TYPES)

if uploaded is not None:
    st.info(f"File: {uploaded.name} • Size: {uploaded.size/1e6:.2f} MB")
    with st.spinner("Preparing audio…"):
        wav_path = tempfile.mktemp(suffix=".wav")
        try:
            extract_audio_to_wav(uploaded.read(), wav_path)
        except RuntimeError as e:
            st.error(str(e))
            st.stop()

    st.success("Audio prepared. Transcribing…")

    # Import here to avoid importing if user never uploads
    from faster_whisper import WhisperModel

    model = load_model(model_size, compute_type)
    segments_out = []
    with st.spinner("Transcribing with Faster-Whisper… this can take a moment for large files"):
        segments, info = model.transcribe(
            wav_path,
            beam_size=beam_size,
            vad_filter=vad_filter,
            language=None,
        )
        for seg in segments:
            segments_out.append(seg)

    try:
        os.unlink(wav_path)
    except OSError:
        pass

    if not segments_out:
        st.error("No speech detected.")
    else:
        st.success(
            f"Done. Detected language: {getattr(info, 'language', 'unknown')}. "
            f"Duration: {getattr(info, 'duration', 'n/a'):.2f}s"
        )
        plaintext = segments_to_plaintext(segments_out)
        srt_text = segments_to_srt(segments_out)
        docx_bytes = write_docx(segments_out)

        st.subheader("Preview")
        st.text_area("Transcript (plain text)", plaintext, height=250)

        st.download_button("⬇️ Download .txt", data=plaintext, file_name="transcript.txt")
        st.download_button("⬇️ Download .srt (subtitles)", data=srt_text, file_name="transcript.srt")
        st.download_button("⬇️ Download .docx (Word)", data=docx_bytes, file_name="transcript.docx")

        st.caption("Tip: .srt is great for syncing captions; .docx includes timecodes per paragraph.")
else:
    st.write("Upload an MP4 (or audio) to get started.")
